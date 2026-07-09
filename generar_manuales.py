# -*- coding: utf-8 -*-
"""Genera los 4 manuales institucionales de ENCL en DOCX y PDF."""

import os
from datetime import datetime

ANIO = datetime.now().year

# ──────────────────────── CONTENIDO DE LOS MANUALES ────────────────────────

MANUALES = {
    'Manual del Instructor': {
        'portada_subtitulo': 'Guía completa para la gestión de aulas digitales',
        'secciones': [
            ('Introducción', [
                'Bienvenido al Manual del Instructor de la Escuela Nacional de Capacitación y Cursos en Línea (ENCL).',
                'Este manual ha sido diseñado para guiarlo paso a paso en el uso de la plataforma educativa, '
                'permitiéndole gestionar sus cursos, crear contenido, evaluar a sus alumnos y emitir certificados '
                'de manera eficiente y profesional.',
                'ENCL es una institución comprometida con la educación de calidad en México. Nuestra plataforma '
                'ofrece herramientas modernas para la enseñanza en línea, combinando tecnología con pedagogía '
                'para brindar la mejor experiencia educativa.'
            ]),
            ('Requisitos técnicos', [
                'Para impartir clases en la plataforma ENCL, asegúrese de contar con:',
                '• Computadora con cámara web y micrófono funcionales.',
                '• Conexión a Internet de al menos 10 Mbps de descarga y 3 Mbps de subida.',
                '• Navegador web actualizado (Google Chrome, Mozilla Firefox o Microsoft Edge).',
                '• Cuenta de correo electrónico activa.',
                '• Opcional: cámara externa, micrófono profesional y tableta gráfica.',
                'Para sesiones en vivo, se recomienda usar Zoom o Google Meet. Asegúrese de tener instalada '
                'la aplicación correspondiente y una cuenta activa.'
            ]),
            ('Acceso a la plataforma', [
                '1. Abra su navegador web y diríjase a la URL de la plataforma.',
                '2. En la página principal, haga clic en "Iniciar sesión".',
                '3. Ingrese su nombre de usuario y contraseña proporcionados por el administrador.',
                '4. Haga clic en el botón "Iniciar sesión".',
                '5. Será redirigido automáticamente al Panel del Instructor.',
                'Si olvida su contraseña, utilice la opción "Recuperar contraseña" en la pantalla de inicio '
                'de sesión y siga las instrucciones.'
            ]),
            ('Configuración del perfil', [
                'Una vez dentro de la plataforma, puede personalizar su perfil:',
                '• Acceda a "Mi Perfil" desde el menú lateral izquierdo.',
                '• Aquí podrá ver su información personal y actualizar su foto de perfil.',
                '• Para cambiar su foto, haga clic en "Examinar", seleccione una imagen y presione "Actualizar foto".',
                '• La foto debe ser cuadrada, preferiblemente de 400x400 píxeles, en formato JPG o PNG.',
                '• Su información de contacto (nombre, email, teléfono) se muestra de forma informativa. '
                'Para modificarla, contacte al administrador del sistema.'
            ]),
            ('Video de bienvenida', [
                'El video de bienvenida es una herramienta poderosa para conectar con sus alumnos:',
                '• Desde "Mi Perfil" encontrará la sección "Video de Bienvenida del Instructor".',
                '• Haga clic en "Seleccionar archivo" y elija un video en formato MP4.',
                '• El video se mostrará en un marco institucional con el logo de ENCL.',
                '• Recomendaciones: duración de 1 a 3 minutos, buena iluminación, audio claro.',
                '• Puede reemplazar el video en cualquier momento subiendo uno nuevo.',
                'El video se reproduce con un diseño profesional que incluye su nombre, el logo institucional '
                'y los colores de la Escuela Nacional de Capacitación.'
            ]),
            ('Administración del Aula Digital', [
                'El Aula Digital es el espacio central de su curso. Para acceder:',
                '• Desde el menú lateral, haga clic en "Mis aulas digitales".',
                '• Verá la lista de sus cursos. Haga clic en "Entrar al aula digital" del curso deseado.',
                '• El aula digital cuenta con pestañas para gestionar: Clases, Alumnos, Exámenes, Horarios y Enlaces.',
                '• Desde aquí puede crear y organizar todo el contenido de su curso.',
                '• Los alumnos inscritos verán el contenido actualizado automáticamente.'
            ]),
            ('Creación de horarios', [
                'Para establecer el horario de sus clases:',
                '• Dentro del Aula Digital, vaya a la pestaña "Horarios".',
                '• Haga clic en "Agregar horario".',
                '• Seleccione el día de la semana, hora de inicio, hora de fin y salón o aula virtual.',
                '• Opcionalmente, puede crear un grupo y asignar el horario a ese grupo específico.',
                '• Los horarios creados se mostrarán en el calendario del alumno.',
                '• Puede crear múltiples horarios para un mismo curso.'
            ]),
            ('Publicación de enlaces (Zoom / Google Meet)', [
                'Para compartir enlaces de videoconferencia con sus alumnos:',
                '• En el Aula Digital, vaya a la pestaña "Enlaces".',
                '• Haga clic en "Agregar enlace".',
                '• Ingrese un título descriptivo (ej. "Clase 1 - Zoom").',
                '• Pegue la URL completa de la reunión de Zoom o Google Meet.',
                '• Seleccione la plataforma correspondiente (Zoom, Meet, Teams, etc.).',
                '• Los alumnos verán los enlaces en la sección "Clases en vivo" de su panel.',
                '• Recuerde actualizar los enlaces si cambian.'
            ]),
            ('Subida de material didáctico', [
                'Para compartir materiales con sus alumnos:',
                '• Dentro del Aula Digital, en la pestaña "Clases", puede agregar material a cada clase.',
                '• Al crear o editar una clase, verá la opción de subir un archivo.',
                '• Formatos aceptados: PDF, Word, PowerPoint, Excel, imágenes y video MP4.',
                '• Los materiales se muestran en la sección "Material" del panel del alumno.',
                '• También puede incluir materiales complementarios como guías de estudio, lecturas y ejercicios.'
            ]),
            ('Publicación de tareas', [
                'Para asignar tareas a sus alumnos:',
                '• Desde el Aula Digital, acceda a la pestaña de gestión correspondiente.',
                '• Cree una nueva tarea especificando título, descripción y fecha de entrega.',
                '• Puede adjuntar un archivo con las instrucciones o el formato de la tarea.',
                '• Los alumnos recibirán la tarea en su panel y podrán subir sus trabajos.',
                '• Usted podrá revisar, calificar y dejar comentarios en cada entrega.'
            ]),
            ('Publicación de clases grabadas', [
                'Para compartir grabaciones de clases:',
                '• En el Aula Digital, pestaña "Clases", seleccione "Grabada" como tipo de clase.',
                '• Ingrese la URL del video (puede ser de YouTube, Vimeo o un archivo subido).',
                '• Proporcione un título descriptivo y una breve descripción del contenido.',
                '• Las clases grabadas aparecen en la sección "Clases grabadas" del alumno.',
                '• Los alumnos pueden verlas en cualquier momento y las veces que deseen.'
            ]),
            ('Creación y administración de exámenes', [
                'Para evaluar a sus alumnos:',
                '• En el Aula Digital, vaya a la pestaña "Exámenes".',
                '• Haga clic en "Crear examen".',
                '• Configure el título, descripción, tiempo límite (en minutos) y calificación mínima aprobatoria.',
                '• Agregue preguntas de opción múltiple (hasta 10 por examen).',
                '• Cada pregunta debe tener 4 opciones, indicando cuál es la correcta.',
                '• Los alumnos reciben su calificación inmediatamente al terminar el examen.',
                '• Puede ver los resultados desde la sección "Calificaciones" del menú principal.'
            ]),
            ('Captura de calificaciones', [
                'Para dar seguimiento al rendimiento de sus alumnos:',
                '• Desde el menú lateral, acceda a "Calificaciones".',
                '• Verá una tabla con todos sus cursos y los alumnos inscritos.',
                '• Las calificaciones de los exámenes se capturan automáticamente.',
                '• Se calcula automáticamente el promedio de cada alumno.',
                '• Puede identificar rápidamente a los alumnos que necesitan apoyo adicional.',
                '• Los colores indican: verde (aprobado) y rojo (reprobado).'
            ]),
            ('Seguimiento del avance de los alumnos', [
                'Para monitorear el progreso de sus alumnos:',
                '• En el Aula Digital, la pestaña "Alumnos" muestra el avance de cada inscrito.',
                '• Puede registrar la asistencia a clases en vivo desde esta misma sección.',
                '• El progreso general del alumno se actualiza automáticamente.',
                '• Use esta información para identificar alumnos en riesgo y ofrecer apoyo oportuno.',
                '• También puede ver el historial completo de cada alumno.'
            ]),
            ('Emisión de certificados', [
                'Los certificados se generan automáticamente cuando un alumno:',
                '• Completa todos los módulos del curso.',
                '• Aprueba todos los exámenes con la calificación mínima requerida.',
                '• Cumple con el porcentaje mínimo de asistencia (si aplica).',
                '• Desde "Certificados" en el menú lateral, puede ver todos los certificados emitidos.',
                '• Cada certificado tiene un folio único y un código QR para verificación.',
                '• Los certificados son válidos y verificables en cualquier momento.'
            ]),
            ('Recomendaciones para clases en línea', [
                '• Prepare su espacio: busque un lugar tranquilo, bien iluminado y sin distracciones.',
                '• Pruebe su equipo: verifique cámara, micrófono y conexión antes de cada sesión.',
                '• Llegue temprano: conéctese 10-15 minutos antes para resolver cualquier incidente técnico.',
                '• Interactúe: fomente la participación activa de los alumnos con preguntas y debates.',
                '• Use materiales visuales: complemente su clase con presentaciones, videos y ejemplos prácticos.',
                '• Sea claro: explique los objetivos de aprendizaje al inicio y resuma los puntos clave al final.',
                '• Grabe las sesiones: permita que los alumnos que no pudieron asistir accedan al contenido.',
                '• Dé seguimiento: revise las tareas y exámenes oportunamente y proporcione retroalimentación.'
            ]),
            ('Preguntas frecuentes', [
                'P: ¿Cómo restablezco mi contraseña?\nR: Use la opción "Recuperar contraseña" en la pantalla de inicio de sesión.',
                'P: ¿Puedo tener más de un curso?\nR: Sí, puede gestionar todos los cursos que le sean asignados.',
                'P: ¿Cómo agrego a un alumno a mi curso?\nR: Los alumnos se inscriben a través de la plataforma. Usted no necesita agregarlos manualmente.',
                'P: ¿Los exámenes se califican automáticamente?\nR: Sí, los exámenes de opción múltiple se califican automáticamente al ser enviados.',
                'P: ¿Puedo editar un examen después de creado?\nR: Actualmente debe crear un nuevo examen con las preguntas modificadas.',
                'P: ¿Cómo veo quién ha visto mi video de bienvenida?\nR: El video de bienvenida está disponible para todos los alumnos inscritos en sus cursos.'
            ]),
            ('Contacto para soporte técnico', [
                'Si requiere asistencia técnica, puede contactarnos a través de:',
                '• Correo electrónico: soporte@encl.edu.mx',
                '• Teléfono: 55 5555 5555',
                '• Formulario de contacto: disponible en la plataforma, sección "Contacto".',
                'Horario de atención: lunes a viernes de 9:00 a 18:00 horas.',
                'Tiempo de respuesta máximo: 24 horas hábiles.'
            ]),
        ]
    },
    'Manual del Alumno': {
        'portada_subtitulo': 'Guía para estudiantes de la plataforma ENCL',
        'secciones': [
            ('Introducción', [
                'Bienvenido a la Escuela Nacional de Capacitación y Cursos en Línea (ENCL).',
                'Este manual le ayudará a navegar la plataforma, inscribirse en cursos, participar en clases '
                'en vivo, realizar exámenes y obtener sus certificados.',
                'ENCL está comprometida con su desarrollo profesional y personal, ofreciendo educación '
                'de calidad accesible para todos los mexicanos.'
            ]),
            ('Requisitos técnicos', [
                '• Computadora, tableta o teléfono inteligente con acceso a Internet.',
                '• Conexión estable de al menos 5 Mbps de descarga.',
                '• Navegador actualizado (Chrome, Firefox, Edge o Safari).',
                '• Cuenta de correo electrónico válida.',
                '• Para clases en vivo: cámara web y micrófono (opcional, pero recomendado).',
            ]),
            ('Creación de cuenta', [
                '1. Vaya a la página principal de la plataforma.',
                '2. Haga clic en "Registrarme" o "Crear cuenta".',
                '3. Complete el formulario con sus datos personales.',
                '4. Elija un nombre de usuario y una contraseña segura.',
                '5. Acepte los términos y condiciones.',
                '6. Haga clic en "Crear cuenta".',
                '7. Una vez registrado, inicie sesión con sus credenciales.',
            ]),
            ('Inicio de sesión', [
                '1. En la página principal, haga clic en "Iniciar sesión".',
                '2. Ingrese su usuario y contraseña.',
                '3. Haga clic en "Iniciar sesión".',
                '4. Será redirigido a su panel de alumno.',
                'Si olvida su contraseña, use la opción "Recuperar contraseña".',
            ]),
            ('Panel del alumno', [
                'Su panel principal muestra:',
                '• Cursos inscritos y progreso.',
                '• Próximas clases en vivo.',
                '• Calendario con fechas importantes.',
                '• Acceso rápido a mensajes y certificados.',
                '• Estadísticas de su avance general.',
            ]),
            ('Inscripción a cursos', [
                '1. Desde el menú, vaya a "Cursos" o "Mis cursos".',
                '2. Explore el catálogo de cursos disponibles.',
                '3. Haga clic en el curso de su interés para ver los detalles.',
                '4. Presione "Inscribirme" para comenzar.',
                '5. El curso aparecerá en su lista de "Mis cursos".',
                '6. Puede inscribirse en múltiples cursos simultáneamente.',
            ]),
            ('Clases en vivo', [
                'Las clases en vivo se imparten a través de Zoom o Google Meet:',
                '• En "Clases en vivo" verá el horario de sus próximas clases.',
                '• Haga clic en el enlace proporcionado para unirse a la sesión.',
                '• Conéctese 5 minutos antes de la hora programada.',
                '• Participe activamente, haga preguntas y tome notas.',
            ]),
            ('Clases grabadas', [
                '• Acceda a "Clases grabadas" para ver las grabaciones disponibles.',
                '• Puede verlas en cualquier momento y las veces que necesite.',
                '• Ideal para repasar contenido o ponerse al día si faltó a una clase.',
            ]),
            ('Material de estudio', [
                '• En "Material" encontrará guías, lecturas y recursos de sus cursos.',
                '• Puede descargar los materiales para consultarlos sin conexión.',
                '• Revise el material antes de cada clase para mejor comprensión.',
            ]),
            ('Tareas', [
                '• Desde "Material" puede ver las tareas asignadas.',
                '• Cada tarea indica su fecha de entrega y las instrucciones.',
                '• Suba su trabajo en el formato solicitado antes de la fecha límite.',
                '• Recibirá retroalimentación y calificación por parte del instructor.',
            ]),
            ('Exámenes', [
                '• En "Exámenes" encontrará las evaluaciones disponibles.',
                '• Lea cuidadosamente cada pregunta antes de responder.',
                '• Recibirá su calificación inmediatamente al finalizar.',
                '• La calificación mínima aprobatoria es 6.0.',
                '• Puede intentar el examen nuevamente si no aprueba.',
            ]),
            ('Certificados', [
                '• Al aprobar un curso, recibirá un certificado digital.',
                '• El certificado incluye un folio único y un código QR.',
                '• Puede descargar e imprimir su certificado desde "Certificados".',
                '• Los empleadores pueden verificar la autenticidad del certificado en línea.',
            ]),
            ('Mensajes', [
                '• Use "Mensajes" para comunicarse con instructores y administradores.',
                '• Recibirá notificaciones de nuevos mensajes en su bandeja.',
                '• Responda oportunamente a los mensajes de sus instructores.',
            ]),
            ('Preguntas frecuentes', [
                'P: ¿Cómo recupero mi contraseña?\nR: Use la opción "Recuperar contraseña" en la pantalla de inicio de sesión.',
                'P: ¿Puedo cambiar mi información personal?\nR: Contacte al administrador para modificar sus datos.',
                'P: ¿Cuánto tiempo tengo para completar un curso?\nR: Depende del curso. Consulte la descripción del curso para más detalles.',
                'P: ¿Los cursos tienen costo?\nR: Algunos cursos son gratuitos y otros tienen costo. Verifique la información del curso.',
            ]),
            ('Contacto para soporte técnico', [
                '• Correo: soporte@encl.edu.mx',
                '• Teléfono: 55 5555 5555',
                '• Horario: lunes a viernes, 9:00 a 18:00 horas.',
            ]),
        ]
    },
    'Manual del Administrador': {
        'portada_subtitulo': 'Guía para la gestión administrativa de la plataforma ENCL',
        'secciones': [
            ('Introducción', [
                'Este manual está dirigido a los administradores de la plataforma ENCL.',
                'Como administrador, usted tiene acceso a herramientas para gestionar usuarios, cursos, '
                'categorías, pagos, becas, certificados y la configuración general del sistema.',
                'Su rol es fundamental para el correcto funcionamiento de la plataforma educativa.'
            ]),
            ('Acceso al panel administrativo', [
                '1. Inicie sesión con su cuenta de administrador.',
                '2. Será redirigido automáticamente al Panel de Administración.',
                '3. Desde aquí puede acceder a todas las secciones de gestión.',
                '4. El menú lateral contiene los enlaces a cada módulo administrativo.',
            ]),
            ('Gestión de usuarios', [
                '• Acceda a "Usuarios" para ver, crear, editar y desactivar usuarios.',
                '• Puede filtrar por rol (alumno, instructor, admin) y estado (activo/inactivo).',
                '• Al crear un usuario, asigne un rol adecuado y una contraseña temporal.',
                '• Puede activar o desactivar usuarios sin eliminarlos.',
                '• Revise periódicamente los usuarios inactivos para mantener la base de datos limpia.',
            ]),
            ('Gestión de cursos', [
                '• En "Cursos" puede crear, editar y eliminar cursos.',
                '• Asigne un instructor responsable a cada curso.',
                '• Configure el nivel, duración, precio y modalidad del curso.',
                '• Puede activar o desactivar cursos según sea necesario.',
                '• Organice los cursos por categorías para facilitar la navegación.',
            ]),
            ('Gestión de categorías', [
                '• Las categorías permiten organizar los cursos por tema.',
                '• Cree categorías como "Tecnología", "Idiomas", "Desarrollo Profesional", etc.',
                '• Cada categoría puede tener un icono representativo.',
                '• Un curso puede pertenecer a una sola categoría.',
            ]),
            ('Gestión de pagos', [
                '• En "Pagos" puede registrar y verificar pagos de los alumnos.',
                '• Los pagos pueden ser en efectivo, transferencia o tarjeta.',
                '• Verifique un pago para activar el acceso del alumno al curso.',
                '• Puede ver el historial completo de pagos de cada alumno.',
                '• Genere reportes de ingresos por periodo.',
            ]),
            ('Gestión de becas', [
                '• Cree programas de becas con diferentes tipos (porcentaje o monto fijo).',
                '• Configure los requisitos, cupo máximo y fechas de vigencia.',
                '• Revise y resuelva las solicitudes de beca de los alumnos.',
                '• Una beca aprobada se aplica automáticamente al costo del curso.',
            ]),
            ('Gestión de certificados', [
                '• En "Certificados" puede ver todos los certificados emitidos en la plataforma.',
                '• Puede validar certificados por folio desde la sección correspondiente.',
                '• Los certificados revocados aparecen como no válidos.',
                '• Cada certificado tiene un código QR único para verificación externa.',
            ]),
            ('Gestión de contenido (CMS)', [
                '• El CMS permite personalizar el contenido de la página principal.',
                '• Puede modificar: título del sitio, textos, misión, visión, colores y datos de contacto.',
                '• Los cambios se reflejan inmediatamente en el sitio público.',
                '• Mantenga la información actualizada para una imagen profesional.',
            ]),
            ('Gestión de imágenes y banners', [
                '• En "Imágenes" puede subir y organizar las imágenes del sitio.',
                '• Los banners son imágenes destacadas que aparecen en la página principal.',
                '• Puede reordenar los banners y establecer su tiempo de visualización.',
                '• Las imágenes deben optimizarse para web (máximo 500 KB).',
            ]),
            ('Reportes', [
                '• La sección "Reportes" le permite visualizar estadísticas clave.',
                '• Consulte: total de alumnos, cursos, ingresos y más.',
                '• Use esta información para la toma de decisiones institucionales.',
            ]),
            ('Preguntas frecuentes', [
                'P: ¿Cómo creo un nuevo instructor?\nR: Desde "Usuarios", cree un nuevo usuario con rol "instructor".',
                'P: ¿Puedo eliminar un curso con alumnos inscritos?\nR: Es recomendable desactivarlo en lugar de eliminarlo.',
                'P: ¿Cómo verifico un pago?\nR: Desde "Pagos", localice el pago pendiente y presione "Verificar".',
            ]),
        ]
    },
    'Manual del Súper Administrador': {
        'portada_subtitulo': 'Guía completa para la administración avanzada de ENCL',
        'secciones': [
            ('Introducción', [
                'El Súper Administrador tiene el nivel más alto de acceso en la plataforma ENCL.',
                'Este manual cubre funciones exclusivas como gestión de patrocinadores, banners, '
                'páginas territoriales y configuración avanzada del sistema.',
                'Se recomienda utilizar este acceso con responsabilidad y solo para tareas administrativas '
                'que requieran privilegios elevados.'
            ]),
            ('Acceso y seguridad', [
                '• Use credenciales seguras y no las comparta con nadie.',
                '• La cuenta de súper administrador no debe usarse para tareas cotidianas.',
                '• Cree cuentas de administrador para la gestión diaria.',
                '• Revise periódicamente los accesos y actividades del sistema.',
                '• Cambie la contraseña regularmente.',
            ]),
            ('Gestión de patrocinadores', [
                '• Los patrocinadores son organizaciones que apoyan los cursos.',
                '• Desde "Patrocinadores" puede crear, editar y eliminar patrocinadores.',
                '• Cada patrocinador puede tener un logo, foto, cargo y mensaje institucional.',
                '• También puede generar un código QR único para cada patrocinador.',
                '• Los patrocinadores se muestran en la página pública de patrocinadores.',
            ]),
            ('Gestión de banners', [
                '• Los banners son elementos visuales destacados en la página principal.',
                '• Puede crear banners con título, subtítulo, imagen y enlace.',
                '• Configure el tiempo de visualización de cada banner en milisegundos.',
                '• Reordene los banners según la prioridad de visualización.',
                '• Active o desactive banners según la campaña vigente.',
            ]),
            ('Páginas territoriales', [
                '• Las páginas territoriales son landing pages para regiones específicas.',
                '• Cada página tiene: nombre, municipio, foto de fondo, mensaje y colores personalizados.',
                '• Puede ajustar la posición y zoom de la imagen de fondo con el editor visual.',
                '• Incluyen código QR, información de contacto y meta tags para SEO.',
                '• La URL de cada página es: /territorial/[slug].',
                '• Use estas páginas para presencia institucional en diferentes regiones.',
            ]),
            ('Configuración avanzada', [
                '• Gestión de favicon, logo y colores institucionales desde el CMS.',
                '• Subida de logo oficial, logo clean (SVG) e imágenes de encabezado.',
                '• Configuración de meta tags para SEO de la página principal.',
                '• Personalización de textos institucionales (misión, visión, about).',
            ]),
            ('Verificación de certificados', [
                '• Todos los certificados emitidos son verificables públicamente.',
                '• La URL de verificación es: /admin/validar-certificado/[folio].',
                '• Los empleadores y terceros pueden verificar la autenticidad sin iniciar sesión.',
                '• Un certificado revocado muestra estado "No válido".',
            ]),
            ('Buenas prácticas', [
                '• Realice respaldos periódicos de la base de datos.',
                '• Mantenga un registro de cambios importantes en la plataforma.',
                '• Capacite a los administradores en el uso correcto del sistema.',
                '• Revise periódicamente los logs y reportes de actividad.',
                '• Actualice la plataforma y sus dependencias regularmente.',
                '• Atienda oportunamente los reportes de soporte técnico.',
            ]),
            ('Preguntas frecuentes', [
                'P: ¿Cómo agrego un nuevo patrocinador?\nR: Desde "Patrocinadores", haga clic en "Nuevo patrocinador" y complete el formulario.',
                'P: ¿Cómo elimino una página territorial?\nR: Desde "Páginas Territoriales", use el botón de eliminar en la página deseada.',
                'P: ¿Puedo duplicar una página territorial?\nR: Sí, use el botón "Duplicar" para crear una copia con los mismos datos.',
                'P: ¿Cómo ajusto la imagen de fondo de una página territorial?\nR: Use el editor Cropper.js integrado en el formulario de edición.',
            ]),
        ]
    }
}

# ──────────────────────── GENERACIÓN DOCX ────────────────────────

def generar_docx(titulo, datos, ruta):
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    import docx.oxml

    doc = Document()

    # Estilos
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)

    # ─── PORTADA ───
    for _ in range(6):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('ESCUELA NACIONAL DE\nCAPACITACIÓN Y CURSOS EN LÍNEA')
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x1a, 0x23, 0x7e)

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(titulo)
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0x0a, 0x0a, 0x1a)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(datos['portada_subtitulo'])
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    for _ in range(4):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'Versión 1.0 — {ANIO}')
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_page_break()

    # ─── ÍNDICE ───
    p = doc.add_paragraph()
    run = p.add_run('Índice')
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x1a, 0x23, 0x7e)

    for i, (titulo_sec, _) in enumerate(datos['secciones'], 1):
        p = doc.add_paragraph(f'{i}. {titulo_sec}')
        p.paragraph_format.space_after = Pt(2)
        run = p.runs[0]
        run.font.size = Pt(11)

    doc.add_page_break()

    # ─── SECCIONES ───
    for i, (titulo_sec, parrafos) in enumerate(datos['secciones'], 1):
        p = doc.add_paragraph()
        run = p.add_run(f'{i}. {titulo_sec}')
        run.bold = True
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0x1a, 0x23, 0x7e)
        p.paragraph_format.space_before = Pt(12)

        for texto in parrafos:
            p = doc.add_paragraph(texto)
            p.paragraph_format.space_after = Pt(4)

    # ─── PIE DE PÁGINA ───
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'ESCUELA NACIONAL DE CAPACITACIÓN Y CURSOS EN LÍNEA — Manual Institucional — {ANIO}')
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.save(ruta)
    print(f'  OK DOCX: {ruta}')


# ──────────────────────── GENERACIÓN PDF ────────────────────────

def generar_pdf(titulo, datos, ruta):
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch, cm
    from reportlab.lib.colors import HexColor
    from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer, PageBreak,
                                     Table, TableStyle, Frame, PageTemplate)
    from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY

    doc = SimpleDocTemplate(
        ruta, pagesize=letter,
        topMargin=2*cm, bottomMargin=2*cm,
        leftMargin=2.5*cm, rightMargin=2.5*cm
    )

    styles = getSampleStyleSheet()
    primary = HexColor('#1a237e')
    dark = HexColor('#0a0a1a')

    title_style = ParagraphStyle('TitleCustom', parent=styles['Title'],
                                  fontSize=24, leading=32, textColor=primary,
                                  spaceAfter=6, alignment=TA_CENTER)

    h1_style = ParagraphStyle('H1Custom', parent=styles['Heading1'],
                               fontSize=18, leading=24, textColor=primary,
                               spaceBefore=14, spaceAfter=8)

    body_style = ParagraphStyle('BodyCustom', parent=styles['Normal'],
                                 fontSize=10, leading=14,
                                 spaceAfter=6, alignment=TA_JUSTIFY)

    bullet_style = ParagraphStyle('BulletCustom', parent=body_style,
                                   leftIndent=20, bulletIndent=8,
                                   spaceAfter=3)

    contents = []

    # ─── PORTADA ───
    for _ in range(8):
        contents.append(Spacer(1, 0.5*inch))
    contents.append(Paragraph('ESCUELA NACIONAL DE<br/>CAPACITACIÓN Y CURSOS EN LÍNEA', title_style))
    contents.append(Spacer(1, 0.3*inch))
    contents.append(Paragraph(f'<b>{titulo}</b>', ParagraphStyle('BigTitle', parent=title_style,
                               fontSize=28, leading=34, textColor=dark)))
    contents.append(Spacer(1, 0.2*inch))
    contents.append(Paragraph(datos['portada_subtitulo'], ParagraphStyle('SubTitle', parent=body_style,
                               fontSize=13, textColor=HexColor('#555555'), alignment=TA_CENTER)))
    for _ in range(5):
        contents.append(Spacer(1, 0.3*inch))
    contents.append(Paragraph(f'Versión 1.0 — {ANIO}', ParagraphStyle('Version', parent=body_style,
                               fontSize=10, textColor=HexColor('#888888'), alignment=TA_CENTER)))
    contents.append(PageBreak())

    # ─── ÍNDICE ───
    contents.append(Paragraph('Índice', h1_style))
    contents.append(Spacer(1, 0.15*inch))
    for i, (titulo_sec, _) in enumerate(datos['secciones'], 1):
        contents.append(Paragraph(f'{i}. {titulo_sec}', body_style))
    contents.append(PageBreak())

    # ─── SECCIONES ───
    for i, (titulo_sec, parrafos) in enumerate(datos['secciones'], 1):
        contents.append(Paragraph(f'{i}. {titulo_sec}', h1_style))
        for texto in parrafos:
            if texto.startswith('•'):
                contents.append(Paragraph(texto, bullet_style))
            elif texto.startswith('P:'):
                contents.append(Paragraph(texto.replace('\n', '<br/>'), body_style))
            else:
                contents.append(Paragraph(texto, body_style))
        contents.append(Spacer(1, 0.1*inch))

    doc.build(contents)
    print(f'  OK PDF:  {ruta}')


# ──────────────────────── EJECUCIÓN ────────────────────────

if __name__ == '__main__':
    output_dir = os.path.join(os.path.dirname(__file__), 'manuales')
    os.makedirs(output_dir, exist_ok=True)

    for nombre, datos in MANUALES.items():
        print(f'\nGenerando {nombre}...')
        nombre_archivo = nombre.lower().replace(' ', '_').replace('ú', 'u').replace('ó', 'o')
        generar_docx(nombre, datos, os.path.join(output_dir, f'{nombre_archivo}.docx'))
        generar_pdf(nombre, datos, os.path.join(output_dir, f'{nombre_archivo}.pdf'))

    print(f'\nOK - Manuales generados en: {output_dir}')
    print('  - Manual del Instructor (DOCX + PDF)')
    print('  - Manual del Alumno (DOCX + PDF)')
    print('  - Manual del Administrador (DOCX + PDF)')
    print('  - Manual del Súper Administrador (DOCX + PDF)')
